# prepare custom widget list
from __future__ import annotations
from ecoscope_workflows_core.decorators import task
from pydantic import BaseModel
from ecoscope_workflows_core.annotations import AnyDataFrame
from typing import Mapping, Hashable, TypeGuard
from typing import Annotated, Any, Literal, Tuple
from ecoscope_workflows_core.tasks.filter._filter import TimeRange
from pydantic import Field
from pydantic.json_schema import SkipJsonSchema


class DocHeadingWidget(BaseModel):
    heading: str | None = None
    level: int = 1


class DocTableWidget(DocHeadingWidget):
    df: AnyDataFrame
    caption: str | None = None


class DocFigureWidget(DocHeadingWidget):
    filepath: str
    caption: str | None = None
    width: float = 5.0  # Default width in inches


DocWidget = DocHeadingWidget | DocTableWidget | DocFigureWidget | list[DocTableWidget] | list[DocFigureWidget]
Predicate = Tuple[str, Literal["=", "!=", "in", "not in"], Any]

WidgetSingle = DocHeadingWidget | DocTableWidget | DocFigureWidget
WidgetOrList = WidgetSingle | list[DocWidget]
WidgetMap = Mapping[Hashable, object]  # values may be widgets or lists/maps (we'll inspect at runtime)
KeyValList = list[tuple[Hashable, object]]


def _is_widget(x: object) -> TypeGuard[WidgetSingle]:
    return isinstance(x, (DocHeadingWidget, DocTableWidget, DocFigureWidget))


def _flatten_values(vals: list[object]) -> list[DocWidget]:
    out: list[DocWidget] = []
    for v in vals:
        if _is_widget(v):
            out.append(v)
        elif isinstance(v, list):
            for item in v:
                if _is_widget(item):
                    out.append(item)
        elif isinstance(v, Mapping):
            out.extend(_flatten_values(list(v.values())))
        elif isinstance(v, tuple) and len(v) == 2 and _is_widget(v[1]):
            out.append(v[1])
        else:
            pass
    return out


class DocGroup(BaseModel):
    """
    First-class grouped doc widget: a group of widgets tied to filter predicates.
    `label` is optional; if omitted we'll derive something sensible from predicates.
    """

    predicates: list[Predicate]
    widgets: list[object]
    label: str | None = None


def _is_group_tuple(x: Any) -> bool:
    # current legacy payloads look like: ((predicates...), widget_or_list)
    return isinstance(x, tuple) and len(x) == 2 and isinstance(x[0], (list, tuple))


def _is_widget_by_type(x: Any) -> bool:
    """
    Check if an object is a widget by examining its class name and attributes.
    This is more robust than isinstance checks when dealing with imports from different modules.
    """
    if hasattr(x, "__class__"):
        class_name = x.__class__.__name__
        if class_name in ("DocHeadingWidget", "DocTableWidget", "DocFigureWidget"):
            return True
        # Also check if it has the expected attributes
        if hasattr(x, "heading") and hasattr(x, "level"):
            if hasattr(x, "filepath"):  # DocFigureWidget
                return True
            elif hasattr(x, "df"):  # DocTableWidget
                return True
            else:  # DocHeadingWidget
                return True
    return False


def _coerce_widget_like(obj: object) -> object:
    # Already a local widget?
    if isinstance(obj, (DocHeadingWidget, DocTableWidget, DocFigureWidget)):
        return obj

    # Cross-module widget? (duck-typed)
    if _is_widget_by_type(obj):
        # Rebuild as local model
        heading = getattr(obj, "heading", None)
        level = getattr(obj, "level", 1)
        caption = getattr(obj, "caption", None)

        if hasattr(obj, "filepath"):
            width = getattr(obj, "width", 5.0)
            return DocFigureWidget(
                heading=heading, level=level, filepath=getattr(obj, "filepath"), caption=caption, width=width
            )
        if hasattr(obj, "df"):
            return DocTableWidget(heading=heading, level=level, df=getattr(obj, "df"), caption=caption)
        # Default to heading if only heading/level are present
        return DocHeadingWidget(heading=heading, level=level)

    # dict → local model
    if isinstance(obj, dict):
        if "filepath" in obj:
            return DocFigureWidget(**obj)  # type: ignore[arg-type]
        if "df" in obj:
            return DocTableWidget(**obj)  # type: ignore[arg-type]
        if "heading" in obj or "level" in obj:
            return DocHeadingWidget(**obj)  # type: ignore[arg-type]
        return obj

    # list of (k,v) → dict → local model
    if isinstance(obj, list) and all(isinstance(t, tuple) and len(t) == 2 for t in obj):
        return _coerce_widget_like(dict(obj))

    return obj


def _is_widget(x: object) -> TypeGuard[WidgetSingle]:
    """
    Type guard that checks if an object is a widget, handling cross-module imports.
    """
    # First try isinstance for local widgets
    if isinstance(x, (DocHeadingWidget, DocTableWidget, DocFigureWidget)):
        return True

    # Then use duck typing for widgets from other modules
    return _is_widget_by_type(x)


def _is_group_tuple(x: Any) -> bool:
    # current legacy payloads look like: ((predicates...), widget_or_list)
    return isinstance(x, tuple) and len(x) == 2 and isinstance(x[0], (list, tuple))


def _coerce_to_docgroup(x: Any) -> DocGroup | Any:
    """
    Accept legacy tuple format and convert to DocGroup, or return individual widgets as-is.
    Legacy examples:
      ((('TemporalGrouper_%B', '=', 'February'),), DocFigureWidget(...))
      ([('field','in',['A','B'])], [DocTableWidget(...), DocFigureWidget(...)])
    """
    # Handle individual widgets directly - use flexible checking
    if _is_widget_by_type(x) or isinstance(x, (DocHeadingWidget, DocTableWidget, DocFigureWidget)):
        return x

    if not _is_group_tuple(x):
        return x

    preds_raw, widgets_raw = x

    # normalize predicates -> list[Predicate]
    preds: list[Predicate] = []
    if isinstance(preds_raw, tuple):
        # either a single predicate triple or a tuple of triples
        if len(preds_raw) == 3 and isinstance(preds_raw[1], str):
            preds = [preds_raw]  # single predicate
        else:
            preds = list(preds_raw)
    elif isinstance(preds_raw, list):
        preds = preds_raw
    else:
        raise TypeError(f"Unsupported predicate structure: {type(preds_raw)}")

    # normalize widgets -> list[widgets]
    if _is_widget_by_type(widgets_raw) or isinstance(widgets_raw, (DocHeadingWidget, DocTableWidget, DocFigureWidget)):
        widgets = [_coerce_widget_like(widgets_raw)]  # type: list[object]
    elif isinstance(widgets_raw, list):
        widgets = []
        for widget in widgets_raw:
            w = _coerce_widget_like(widget)
            if _is_widget_by_type(w) or isinstance(w, (DocHeadingWidget, DocTableWidget, DocFigureWidget)):
                widgets.append(w)
            # else: ignore unrecognized members gracefully
    else:
        # final fallback: try to coerce once
        coerced = _coerce_widget_like(widgets_raw)
        if _is_widget_by_type(coerced) or isinstance(coerced, (DocHeadingWidget, DocTableWidget, DocFigureWidget)):
            widgets = [coerced]
        else:
            raise TypeError(f"Unsupported widget structure in group: {type(widgets_raw)}")

    # try to derive a friendly label, e.g., "February" when you group by month
    derived_label = None
    if len(preds) == 1:
        field, op, val = preds[0]
        if isinstance(val, (str, int)):
            derived_label = str(val)

    return DocGroup(predicates=preds, widgets=widgets, label=derived_label)


def _flatten_doc_items(items: Any) -> list[Any]:
    """
    Flatten nested lists/tuples but do NOT break widget objects.
    We'll handle tuple->DocGroup coercion separately.
    """
    out: list[Any] = []
    if isinstance(items, list):
        for x in items:
            out.extend(_flatten_doc_items(x))
    else:
        out.append(items)
    return out


@task
def prepare_widget_list(
    widgets: Annotated[object, Field(description="Widget(s) or containers of widgets")],
) -> list[object]:
    """
    Normalize widgets into a flat list[DocWidget].
    Accepts:
      - a single widget (local/cross-module)
      - a list of widgets
      - a mapping: key -> widget or list-of-widgets
      - a list of (key, value) pairs
    """

    def _flatten_values(vals: list[object]) -> list[object]:
        out: list[object] = []
        print("DEBUG _flatten_values called with:", type(vals), "len:", len(vals))
        for i, v in enumerate(vals):
            print(f"  DEBUG item[{i}] type:", type(v), "value:", v)
            v2 = _coerce_widget_like(v)
            print("   DEBUG after _coerce_widget_like:", type(v2), v2)
            if _is_widget_by_type(v2) or isinstance(v2, (DocHeadingWidget, DocTableWidget, DocFigureWidget)):
                print("   DEBUG recognized widget:", v2)
                out.append(v2)
            elif isinstance(v2, list):
                print("   DEBUG nested list found, recursing")
                for j, item in enumerate(v2):
                    item2 = _coerce_widget_like(item)
                    print(f"    DEBUG nested[{j}] after coerce:", type(item2), item2)
                    if _is_widget_by_type(item2) or isinstance(
                        item2, (DocHeadingWidget, DocTableWidget, DocFigureWidget)
                    ):
                        out.append(item2)
            elif isinstance(v2, Mapping):
                print("   DEBUG nested mapping found, recursing")
                out.extend(_flatten_values(list(v2.values())))
        print("DEBUG _flatten_values returning len:", len(out))
        return out

    print("DEBUG prepare_widget_list called with:", type(widgets), "value:", widgets)

    if isinstance(widgets, list) and widgets:
        print("DEBUG widgets is list, first item type:", type(widgets[0]), "value:", widgets[0])

    # single widget
    if _is_widget_by_type(widgets) or isinstance(widgets, (DocHeadingWidget, DocTableWidget, DocFigureWidget)):
        print("DEBUG single widget case hit")
        return [widgets]

    # mapping
    if isinstance(widgets, Mapping):
        print("DEBUG mapping case hit")
        return _flatten_values(list(widgets.values()))

    # list / iterable
    if isinstance(widgets, list):
        print("DEBUG list case hit")
        if widgets and isinstance(widgets[0], tuple) and len(widgets[0]) == 2:
            print("DEBUG list of (key, value) tuples case hit")
            return _flatten_values([dict(widgets)])
        return _flatten_values(widgets)

    # last-ditch
    print("DEBUG last-ditch coercion case hit")
    coerced = _coerce_widget_like(widgets)
    print("DEBUG coerced type:", type(coerced), "value:", coerced)
    if _is_widget_by_type(coerced) or isinstance(coerced, (DocHeadingWidget, DocTableWidget, DocFigureWidget)):
        print("DEBUG last-ditch recognized widget")
        return [coerced]

    print("DEBUG returning empty list")
    return []


def add_table(doc, table_widget, table_index):
    if table_widget.heading:
        doc.add_heading(table_widget.heading, level=table_widget.level)
    df = table_widget.df
    table = doc.add_table(rows=len(df.index) + 1, cols=len(df.columns) + 1)
    header_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        header_cells[i + 1].text = str(column)

    for i, idx in enumerate(df.index):
        row_cells = table.rows[i + 1].cells
        # Add index
        row_cells[0].text = str(idx)

        # Add row data
        for j, value in enumerate(df.iloc[i]):
            row_cells[j + 1].text = str(value)

    doc.add_paragraph(
        f"Table {table_index}: {table_widget.caption}" if table_widget.caption else f"Table {table_index}"
    )


def add_figure(doc, widget, index):
    from docx.shared import Inches
    import os

    path = widget.filepath
    print(f"DEBUG add_figure: START - processing widget {index}")
    print(f"DEBUG add_figure: path = '{path}'")
    print(f"DEBUG add_figure: widget heading = '{widget.heading}'")
    print(f"DEBUG add_figure: widget width = {widget.width}")

    # Your paths are already clean, so minimal processing needed
    if isinstance(path, str) and path.startswith("file://"):
        path = path[7:]
        print(f"DEBUG add_figure: removed file:// prefix, new path = '{path}'")

    # File existence and readability checks
    print("DEBUG add_figure: checking if path exists...")
    if not os.path.exists(path):
        print(f"ERROR add_figure: FILE DOES NOT EXIST -> '{path}'")
        # List directory contents to see what's actually there
        dir_path = os.path.dirname(path)
        if os.path.exists(dir_path):
            files = os.listdir(dir_path)
            print(f"DEBUG add_figure: directory contents: {files}")
        else:
            print(f"ERROR add_figure: directory doesn't exist: '{dir_path}'")
        return

    print("DEBUG add_figure: file exists OK")

    if not os.access(path, os.R_OK):
        print(f"ERROR add_figure: file is not readable -> '{path}'")
        return

    print("DEBUG add_figure: file is readable OK")

    # Get file info
    try:
        file_size = os.path.getsize(path)
        print(f"DEBUG add_figure: file size = {file_size} bytes")
    except Exception as e:
        print(f"ERROR add_figure: could not get file size: {e}")
        return

    # Add heading if present
    if widget.heading:
        print(f"DEBUG add_figure: adding heading '{widget.heading}' at level {widget.level}")
        doc.add_heading(widget.heading, level=widget.level)
        print("DEBUG add_figure: heading added successfully")
    else:
        print("DEBUG add_figure: no heading to add")

    # Try to add the picture
    print("DEBUG add_figure: attempting to add picture...")
    try:
        print(f"DEBUG add_figure: calling doc.add_picture with width={widget.width} inches")
        picture = doc.add_picture(path, width=Inches(widget.width))
        print(f"DEBUG add_figure: doc.add_picture returned: {type(picture)}")
        print("DEBUG add_figure: picture added successfully OK")

    except Exception as e:
        print(f"ERROR add_figure: failed to add picture '{path}': {type(e).__name__}: {e}")
        print(f"ERROR add_figure: full exception: {str(e)}")
        import traceback

        traceback.print_exc()
        return

    # Add caption
    caption_text = f"Figure {index}: {widget.caption}" if widget.caption else f"Figure {index}"
    print(f"DEBUG add_figure: adding caption: '{caption_text}'")


@task
def gather_document(
    title: Annotated[str, Field(description="The document title")],
    time_range: Annotated[TimeRange | SkipJsonSchema[None], Field(description="Time range filter")],
    # Accept anything, then coerce to supported types inside the function.
    doc_widgets: Annotated[
        list[object], Field(description="List of document components to gather (widgets or grouped widgets)")
    ],
    root_path: Annotated[str, Field(description="Root path to persist text to")],
    filename: Annotated[
        str, Field(description="The filename to save the document as, without extension. The extension will be .docx")
    ],
    logo_path: Annotated[str | SkipJsonSchema[None], Field(description="The logo file path")] = None,
) -> Annotated[str, Field(description="The saved file path")]:
    import os
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    print("DEBUG gather_document called")
    print(f"  title={title}")
    print(f"  time_range={time_range}")
    print(f"  doc_widgets type={type(doc_widgets)}, len={len(doc_widgets)}")
    print(f"  root_path={root_path}, filename={filename}, logo_path={logo_path}")

    doc = Document()

    # --- Title
    heading1 = doc.add_heading(title, level=1)
    heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.space_before = Inches(2)
    for run in heading1.runs:
        run.font.size = Pt(30)

    # --- Time range (optional)
    if time_range:
        fmt = time_range.time_format
        formatted_time_range = f"From {time_range.since.strftime(fmt)} to {time_range.until.strftime(fmt)}"
        heading2 = doc.add_heading(formatted_time_range, level=2)
        heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading2.runs:
            run.font.size = Pt(15)
        print("DEBUG added time_range heading:", formatted_time_range)
    else:
        print("DEBUG no time_range provided")

    # --- Logo (optional)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Inches(3)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    if logo_path:
        print("DEBUG adding logo:", logo_path)
        run.add_picture(logo_path, width=Inches(1.5))
    else:
        print("DEBUG no logo provided")

    doc.add_page_break()
    print("DEBUG added title page + page break")

    # ---- Normalize inputs
    normalized = [_coerce_to_docgroup(x) for x in _flatten_doc_items(doc_widgets)]
    print("DEBUG normalized widgets:", normalized)

    table_index = 0
    figure_index = 0

    for idx, item in enumerate(normalized):
        print(f"DEBUG processing item[{idx}] type={type(item)} value={item}")
        item = _coerce_widget_like(item)
        print("  DEBUG after _coerce_widget_like:", type(item), item)

        # 1) Group
        if isinstance(item, DocGroup):
            print("  DEBUG DocGroup found with label:", item.label)
            if item.label:
                doc.add_heading(str(item.label), level=2)
            for j, w in enumerate(_flatten_doc_items(item.widgets)):
                print(f"    DEBUG group widget[{j}] type={type(w)} value={w}")
                # --- CHECK MOST SPECIFIC SUBCLASSES FIRST ---
                if isinstance(w, DocTableWidget):
                    table_index += 1
                    print("     DEBUG adding table index:", table_index, "caption:", w.caption)
                    add_table(doc, w, table_index)
                elif isinstance(w, DocFigureWidget):
                    figure_index += 1
                    print("     DEBUG adding figure index:", figure_index, "caption:", w.caption)
                    add_figure(doc, w, figure_index)
                elif isinstance(w, DocHeadingWidget):  # This should be LAST since it's the base class
                    print("     DEBUG adding heading:", w.heading)
                    doc.add_heading(w.heading, level=w.level)

        # 2) Plain widgets - CHECK MOST SPECIFIC SUBCLASSES FIRST
        elif isinstance(item, DocTableWidget):  # Check DocTableWidget before DocHeadingWidget
            table_index += 1
            print("  DEBUG adding plain table index:", table_index, "caption:", item.caption)
            add_table(doc, item, table_index)
        elif isinstance(item, DocFigureWidget):  # Check DocFigureWidget before DocHeadingWidget
            figure_index += 1
            print("  DEBUG adding plain figure index:", figure_index, "caption:", item.caption)
            add_figure(doc, item, figure_index)
        elif isinstance(item, DocHeadingWidget):  # This should be LAST since it's the base class
            print("  DEBUG adding plain heading:", item.heading)
            doc.add_heading(item.heading, level=item.level)

        # 3) Legacy tuple
        elif _is_group_tuple(item):
            print("  DEBUG legacy group tuple found:", item)
            coerced = _coerce_to_docgroup(item)
            if isinstance(coerced, DocGroup):
                print("   DEBUG coerced legacy tuple into DocGroup with label:", coerced.label)
                if coerced.label:
                    doc.add_heading(str(coerced.label), level=2)
                for j, w in enumerate(_flatten_doc_items(coerced.widgets)):
                    print(f"    DEBUG legacy group widget[{j}] type={type(w)} value={w}")
                    # --- CHECK MOST SPECIFIC SUBCLASSES FIRST ---
                    if isinstance(w, DocTableWidget):
                        table_index += 1
                        print("     DEBUG adding table index:", table_index, "caption:", w.caption)
                        add_table(doc, w, table_index)
                    elif isinstance(w, DocFigureWidget):
                        figure_index += 1
                        print("     DEBUG adding figure index:", figure_index, "caption:", w.caption)
                        add_figure(doc, w, figure_index)
                    elif isinstance(w, DocHeadingWidget):  # This should be LAST
                        print("     DEBUG adding heading:", w.heading)
                        doc.add_heading(w.heading, level=w.level)

        else:
            print("  DEBUG skipping unknown item:", item)
            continue

    if root_path.startswith("file://"):
        print("DEBUG stripping file:// prefix from root_path")
        root_path = root_path[7:]

    path = os.path.join(root_path, f"{filename}.docx")
    print("DEBUG saving document to:", path)
    doc.save(path)
    print("DEBUG document saved successfully")
    return str(path)
