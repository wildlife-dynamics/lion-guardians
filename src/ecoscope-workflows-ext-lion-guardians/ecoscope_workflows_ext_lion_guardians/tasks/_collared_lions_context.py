import os
import uuid
import pandas as pd
from pathlib import Path
from pydantic import Field
from datetime import datetime
from docxtpl import DocxTemplate
from ecoscope_workflows_core.decorators import task
from typing import Annotated, Optional, Dict, Any
from ecoscope_workflows_core.tasks.filter._filter import TimeRange
from ecoscope_workflows_ext_custom.tasks.io._path_utils import remove_file_scheme
from ecoscope_workflows_core.annotations import AnyDataFrame
from ecoscope_workflows_ext_ste.tasks._mapbook_context import _format_temporal_grouper


@task
def create_context_page_lg(
    template_path: Annotated[
        str,
        Field(
            description="Path to the .docx template file.",
        ),
    ],
    output_dir: Annotated[
        str,
        Field(
            description="Directory to save the generated .docx file.",
        ),
    ],
    context: Annotated[
        dict,
        Field(
            description="Dictionary with context values for the template.",
        ),
    ],
    filename: Annotated[
        Optional[str],
        Field(
            description="Optional filename . If not provided, a random UUID-based filename will be generated.",
            exclude=True,
        ),
    ] = None,
) -> Annotated[
    str,
    Field(
        description="Full path to the generated .docx file.",
    ),
]:
    """
    Create a context page document from a template and context dictionary.

    Args:
        template_path (str): Path to the .docx template file.
        output_dir (str): Directory to save the generated .docx file.
        context (dict): Dictionary with context values for the template.
        filename (str, optional): Optional filename for the generated file.
            If not provided, a random UUID-based filename will be generated.

    Returns:
        str: Full path to the generated .docx file.
    """
    # Normalize paths
    template_path = remove_file_scheme(template_path)
    output_dir = remove_file_scheme(output_dir)

    # Validate paths
    if not template_path.strip():
        raise ValueError("template_path is empty after normalization")
    if not output_dir.strip():
        raise ValueError("output_dir is empty after normalization")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        filename = "context_page_.docx"
    output_path = Path(output_dir) / filename

    doc = DocxTemplate(template_path)
    doc.render(context)
    doc.save(output_path)
    return str(output_path)


@task
def create_cl_ctx_cover(
    count: int,
    report_period: TimeRange,
    prepared_by: str,
) -> Dict[str, str]:
    """
    Build a dictionary with the mapbook report template values.

    Args:
        count (int): Total number of subjects or records.
        report_period (TimeRange): Object with 'since', 'until', and 'time_format' attributes.
        prepared_by (str): Name of the person or organization preparing the report.

    Returns:
        Dict[str, str]: Structured dictionary with formatted metadata.
    """

    formatted_date = datetime.now()
    formatted_date_str = formatted_date.strftime("%Y-%m-%d %H:%M:%S")
    fmt = getattr(report_period, "time_format", "%Y-%m-%d")
    formatted_time_range = f"{report_period.since.strftime(fmt)} to {report_period.until.strftime(fmt)}"

    # Return structured dictionary
    return {
        "report_id": f"REP-{uuid.uuid4().hex[:8].upper()}",
        "subject_count": str(count),
        "time_generated": formatted_date_str,
        "report_period": formatted_time_range,
        "prepared_by": prepared_by,
    }


def validate_image_path(field_name: str, path: str) -> None:
    """Validate that an image file exists and has valid extension."""
    normalized_path = remove_file_scheme(path)

    if not os.path.exists(normalized_path):
        raise FileNotFoundError(f"Image file for '{field_name}' not found: {normalized_path}")

    valid_extensions = {".png", ".jpg", ".jpeg"}
    if Path(normalized_path).suffix.lower() not in valid_extensions:
        raise ValueError(
            f"Invalid image format for '{field_name}': {Path(normalized_path).suffix}. "
            f"Expected one of {valid_extensions}"
        )

    print(f" Validated image for '{field_name}': {normalized_path}")


@task
def create_collared_lions_grouper_ctx(
    grouper_name: tuple | list | str | None,
    df: AnyDataFrame,
    total_distance: int | float | None,
    home_range: str | None,
    speed_map: str | None,
) -> Dict[str, str | int | float | None]:
    """
    Create context dictionary for mapbook with grouper information and map paths.

    Args:
        grouper_name: The grouper identifier (can be various types)
        df: The dataframe to extract grouper values from
        total_distance: Total distance value for the group
        home_range: Path to the home range map image
        speed_map: Path to the speed map image

    Returns:
        Dictionary containing grouper_value, total_distance, home_range_map, and speed_map
    """

    # Extract grouper value dynamically
    grouper_value = "All"
    print(f"grouper name raw: {grouper_name} (type: {type(grouper_name)})")

    if grouper_name:
        if isinstance(grouper_name, str):
            grouper_value = grouper_name
        elif isinstance(grouper_name, (list, tuple)) and len(grouper_name) > 0:
            # Check if it's a tuple structure like (('index_name', 'All'),)
            first_item = grouper_name[0]

            # Handle tuple structure (('index_name', 'All'),)
            if isinstance(first_item, tuple) and len(first_item) == 2:
                key, value = first_item
                if key == "index_name" and value == "All":
                    grouper_value = "All"
                else:
                    grouper_value = str(value)
                print(f"Extracted from tuple structure: {grouper_value}")
            # Handle grouper objects
            elif hasattr(first_item, "__class__"):
                grouper = first_item
                grouper_type = grouper.__class__.__name__
                print(f"grouper_type: {grouper_type}")

                if grouper_type == "ValueGrouper":
                    index_name = getattr(grouper, "index_name", None)
                    if df is not None and index_name and index_name in df.columns:
                        unique_values = df[index_name].unique()
                        if len(unique_values) == 1:
                            grouper_value = str(unique_values[0])
                        else:
                            grouper_value = index_name
                    else:
                        grouper_value = index_name if index_name else "Value"

                elif grouper_type == "TemporalGrouper":
                    grouper_value = _format_temporal_grouper(grouper, df)

                elif grouper_type == "AllGrouper":
                    grouper_value = "All"
                else:
                    grouper_value = str(grouper_name)
            else:
                grouper_value = str(first_item)
        else:
            grouper_value = str(grouper_name)

    print(f"grouper_value: {grouper_value}")

    # Build context with the required keys
    ctx = {
        "grouper_value": grouper_value,
        "total_distance": total_distance,
        "home_range_map": home_range,
        "speed_map": speed_map,
    }

    print(f"Context: {ctx}")
    return ctx


@task
def merge_cl_files(
    cover_page_path: Annotated[str, Field(description="Path to the cover page .docx file")],
    context_page_items: Annotated[list[Any], Field(description="List of context page document paths to merge.")],
    output_dir: Annotated[str, Field(description="Directory where combined docx will be written")],
    filename: Annotated[Optional[str], Field(description="Optional output filename")] = None,
) -> Annotated[str, Field(description="Path to the combined .docx file")]:
    """
    Combine cover + context pages into a single DOCX.
    Orders context pages by calendar month if month names are detected.
    """
    import os
    import calendar
    from pathlib import Path
    from docx import Document
    from docxcompose.composer import Composer

    # Build month lookup from calendar module
    MONTH_LOOKUP = {name.lower(): idx for idx, name in enumerate(calendar.month_name) if name}
    MONTH_ABBR_LOOKUP = {name.lower(): idx for idx, name in enumerate(calendar.month_abbr) if name}

    def is_skip_sentinel(obj):
        """Check if object is a SkipSentinel."""
        return hasattr(obj, "__class__") and "SkipSentinel" in obj.__class__.__name__

    def extract_path(item):
        """Extract a valid file path from various item formats."""
        # Check for SkipSentinel
        if is_skip_sentinel(item):
            return None

        if isinstance(item, str):
            return item

        if isinstance(item, (list, tuple)):
            # Check for SkipSentinel in tuples
            for x in item:
                if is_skip_sentinel(x):
                    return None
                if isinstance(x, str) and os.path.exists(x):
                    return x
            # Fallback: return first string found
            for x in reversed(item):
                if isinstance(x, str):
                    return x
            return None

        return None

    def detect_month(path: str):
        """Detect month name or abbreviation in filename/path."""
        name = os.path.basename(path).lower()
        for month, idx in MONTH_LOOKUP.items():
            if month in name:
                return idx
        for abbr, idx in MONTH_ABBR_LOOKUP.items():
            if abbr in name:
                return idx
        return None

    # Normalize paths - filter out None and SkipSentinel
    normalized_paths = []
    print(f"Context page items: {context_page_items}")
    for idx, item in enumerate(context_page_items):
        if item is None or is_skip_sentinel(item):
            print(f"Skipping item {idx}: None or SkipSentinel")
            continue

        path = extract_path(item)
        if path is not None:
            normalized_paths.append(path)
        else:
            print(f"Skipping item {idx}: Could not extract valid path")

    if not normalized_paths:
        print("Warning: No valid context pages to merge, returning cover page only")
        # Just save cover page as output
        output_dir = remove_file_scheme(output_dir)
        os.makedirs(output_dir, exist_ok=True)
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"overall_report_{timestamp}.docx"
        output_path = Path(output_dir) / filename

        master = Document(cover_page_path)
        master.save(str(output_path))
        return str(output_path)

    if not os.path.exists(cover_page_path):
        raise FileNotFoundError(f"Cover page file not found: {cover_page_path}")

    for p in normalized_paths:
        if not os.path.exists(p):
            raise FileNotFoundError(f"Context page file not found: {p}")

    # Calendar-aware ordering
    with_month = []
    without_month = []
    for i, path in enumerate(normalized_paths):
        month_idx = detect_month(path)
        if month_idx is not None:
            with_month.append((month_idx, path))
        else:
            without_month.append((i, path))

    with_month.sort(key=lambda x: x[0])
    ordered_paths = [p for _, p in with_month] + [p for _, p in without_month]

    output_dir = remove_file_scheme(output_dir)
    if not output_dir.strip():
        raise ValueError("output_dir is empty after normalization")
    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"overall_report_{timestamp}.docx"

    output_path = Path(output_dir) / filename

    master = Document(cover_page_path)
    composer = Composer(master)

    for doc_path in ordered_paths:
        doc = Document(doc_path)
        composer.append(doc)

    composer.save(str(output_path))
    return str(output_path)


@task
def create_guardians_ctx_cover(
    report_period: TimeRange,
    prepared_by: str,
) -> Dict[str, str]:
    """
    Build a dictionary with the mapbook report template values.

    Args:
        report_period (TimeRange): Object with 'since', 'until', and 'time_format' attributes.
        prepared_by (str): Name of the person or organization preparing the report.

    Returns:
        Dict[str, str]: Structured dictionary with formatted metadata.
    """

    formatted_date = datetime.now()
    formatted_date_str = formatted_date.strftime("%Y-%m-%d %H:%M:%S")
    fmt = getattr(report_period, "time_format", "%Y-%m-%d")
    formatted_time_range = f"{report_period.since.strftime(fmt)} to {report_period.until.strftime(fmt)}"

    # Return structured dictionary
    return {
        "report_id": f"REP-{uuid.uuid4().hex[:8].upper()}",
        "time_generated": formatted_date_str,
        "report_period": formatted_time_range,
        "prepared_by": prepared_by,
    }


@task
def create_guardians_grouper_ctx(
    grouper_name: tuple | list | str | None,
    df: AnyDataFrame,
    total_patrols: int | float | None,
    total_distance: int | float | None,
    total_time: int | float | None,
    patrol_events_track_map: str | None,
    patrol_time_density_map: str | None,
    events_pie_chart: str | None,
    events_time_series_bar_chart: str | None,
    patrol_events: str | None,
    event_efforts: str | None,
    month_stats: str | None,
    guardian_stats: str | None,
) -> Dict[str, str | int | float | None | list]:
    """
    Create context dictionary for guardians report with grouper information and map paths.

    Args:
        grouper_name: The grouper identifier (can be various types)
        df: The dataframe to extract grouper values from
        total_patrols: Total number of patrols
        total_distance: Total distance value for the group (in km)
        total_time: Total time value for the group (in seconds)
        patrol_events_track_map: Path to the patrol events track map image
        patrol_time_density_map: Path to the time density map image
        events_pie_chart: Path to the events pie chart image
        events_time_series_bar_chart: Path to the events time series bar chart image
        patrol_events: Path to patrol events CSV file
        event_efforts: Path to event efforts CSV file
        month_stats: Path to month statistics CSV file
        guardian_stats: Path to guardian statistics CSV file

    Returns:
        Dictionary containing all context data for the guardians report template
    """

    # Extract grouper value dynamically
    grouper_value = "All"
    print(f"grouper name raw: {grouper_name} (type: {type(grouper_name)})")

    if grouper_name:
        if isinstance(grouper_name, str):
            grouper_value = grouper_name
        elif isinstance(grouper_name, (list, tuple)) and len(grouper_name) > 0:
            # Check if it's a tuple structure like (('index_name', 'All'),)
            first_item = grouper_name[0]

            # Handle tuple structure (('index_name', 'All'),)
            if isinstance(first_item, tuple) and len(first_item) == 2:
                key, value = first_item
                if key == "index_name" and value == "All":
                    grouper_value = "All"
                else:
                    grouper_value = str(value)
                print(f"Extracted from tuple structure: {grouper_value}")
            # Handle grouper objects
            elif hasattr(first_item, "__class__"):
                grouper = first_item
                grouper_type = grouper.__class__.__name__
                print(f"grouper_type: {grouper_type}")

                if grouper_type == "ValueGrouper":
                    index_name = getattr(grouper, "index_name", None)
                    if df is not None and index_name and index_name in df.columns:
                        unique_values = df[index_name].unique()
                        if len(unique_values) == 1:
                            grouper_value = str(unique_values[0])
                        else:
                            grouper_value = index_name
                    else:
                        grouper_value = index_name if index_name else "Value"

                elif grouper_type == "TemporalGrouper":
                    grouper_value = _format_temporal_grouper(grouper, df)

                elif grouper_type == "AllGrouper":
                    grouper_value = "All"
                else:
                    grouper_value = str(grouper_name)
            else:
                grouper_value = str(first_item)
        else:
            grouper_value = str(grouper_name)

    print(f"grouper_value: {grouper_value}")

    # Safely convert total_time from seconds to hours
    total_time_hours = None
    if total_time is not None:
        total_time_hours = round(total_time / 3600, 1)

    # Safely read CSV files and convert to list of dictionaries
    def safe_read_csv(file_path: str | None) -> list:
        """Safely read CSV file and return list of dicts, or empty list if file is None/invalid."""
        if file_path is None:
            print("Warning: CSV file path is None")
            return []
        try:
            return pd.read_csv(file_path).to_dict(orient="records")
        except Exception as e:
            print(f"Error reading CSV file {file_path}: {e}")
            return []

    patrol_events_df = safe_read_csv(patrol_events)
    event_efforts_df = safe_read_csv(event_efforts)
    month_stats_df = safe_read_csv(month_stats)
    guardians_stats_df = safe_read_csv(guardian_stats)

    # Build context with the required keys
    ctx = {
        "grouper_value": grouper_value,
        "total_patrols": str(total_patrols) if total_patrols is not None else "0",
        "total_distance": str(total_distance) if total_distance is not None else "0",
        "total_time": str(total_time_hours) if total_time_hours is not None else "0",
        "patrol_events_track_map": patrol_events_track_map,
        "patrol_time_density_map": patrol_time_density_map,
        "events_pie_chart": events_pie_chart,
        "events_time_series_bar_chart": events_time_series_bar_chart,
        "patrol_events": patrol_events_df,
        "event_efforts": event_efforts_df,
        "month_stats": month_stats_df,
        "guardian_stats": guardians_stats_df,
    }

    print(f"Guardians Context: {ctx}")
    return ctx
